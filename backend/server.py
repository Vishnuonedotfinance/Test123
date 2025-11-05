from fastapi import FastAPI, APIRouter, HTTPException, Depends, BackgroundTasks, Response, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Literal
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from mailmerge import MailMerge
import random
import shutil
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'piperocket-secret-key-2025')
JWT_ALGORITHM = 'HS256'

app = FastAPI()
api_router = APIRouter(prefix="/api")
security = HTTPBearer()

# ============= MODELS =============

class UserRole(BaseModel):
    role: Literal['Admin', 'Director', 'Staff']

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    email: EmailStr
    mobile: Optional[str] = None
    role: Literal['Admin', 'Director', 'Staff']
    status: Literal['Active', 'Invited'] = 'Active'
    password_hash: Optional[str] = None
    otp_verified: bool = False
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class UserCreate(BaseModel):
    name: str
    email: EmailStr
    mobile: Optional[str] = None
    role: Literal['Admin', 'Director', 'Staff']
    password: str

class LoginRequest(BaseModel):
    email: EmailStr
    password: str

class OTPVerifyRequest(BaseModel):
    email: EmailStr
    otp: str

class Client(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"client_{uuid.uuid4().hex[:8]}")
    org: str = "piperocket"
    client_name: str
    address: str
    start_date: str
    tenure_months: int
    end_date: str = ""
    currency_preference: Literal['USD', 'INR'] = 'INR'
    service: Literal['PPC', 'SEO', 'Content', 'Backlink']
    amount_inr: float
    amount_ppc: Optional[float] = None
    amount_seo: Optional[float] = None
    authorised_signatory: str
    signatory_designation: str
    gst: str
    poc_name: str
    poc_email: EmailStr
    poc_designation: str
    poc_mobile: str
    approver_user_id: str
    sign_status: Literal['Signed', 'Not signed'] = 'Not signed'
    client_status: Literal['Active', 'Churned'] = 'Active'
    agreement_status: Literal['Live', 'Expired'] = 'Live'
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ClientCreate(BaseModel):
    client_name: str
    address: str
    start_date: str
    tenure_months: int
    currency_preference: Literal['USD', 'INR'] = 'INR'
    service: Literal['PPC', 'SEO', 'Content', 'Backlink']
    amount_inr: float
    amount_ppc: Optional[float] = None
    amount_seo: Optional[float] = None
    authorised_signatory: str
    signatory_designation: str
    gst: str
    poc_name: str
    poc_email: EmailStr
    poc_designation: str
    poc_mobile: str
    approver_user_id: str

class Contractor(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"contractor_{uuid.uuid4().hex[:8]}")
    name: str
    doj: str
    start_date: str
    tenure_months: int
    end_date: str = ""
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    pan: str
    aadhar: str
    mobile: str
    personal_email: EmailStr
    bank_name: str
    account_holder: str
    account_no: str
    ifsc: str
    address_1: str
    pincode: str
    city: str
    address_2: Optional[str] = None
    department: Literal['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']
    projects: List[str] = Field(default_factory=list)
    monthly_retainer_inr: float
    designation: str
    approver_user_id: str
    sign_status: Literal['Signed', 'Not signed'] = 'Not signed'
    status: Literal['Active', 'Terminated'] = 'Active'
    agreement_status: Literal['Live', 'Expired'] = 'Live'
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ContractorCreate(BaseModel):
    name: str
    doj: str
    start_date: str
    tenure_months: int
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    pan: str
    aadhar: str
    mobile: str
    personal_email: EmailStr
    bank_name: str
    account_holder: str
    account_no: str
    ifsc: str
    address_1: str
    pincode: str
    city: str
    address_2: Optional[str] = None
    department: Literal['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']
    projects: List[str] = Field(default_factory=list)
    monthly_retainer_inr: float
    designation: str
    approver_user_id: str

class Employee(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"emp_{uuid.uuid4().hex[:8]}")
    doj: str
    work_email: EmailStr
    emp_id: str
    first_name: str
    last_name: str
    father_name: str
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    mobile: str
    personal_email: EmailStr
    pan: str
    aadhar: str
    uan: str
    pf_account_no: str
    bank_name: str
    account_no: str
    ifsc: str
    branch: str
    address: str
    pincode: str
    city: str
    monthly_gross_inr: float
    department: Literal['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']
    projects: List[str] = Field(default_factory=list)
    approver_user_id: str
    status: Literal['Active', 'Terminated'] = 'Active'
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class EmployeeCreate(BaseModel):
    doj: str
    work_email: EmailStr
    emp_id: str
    first_name: str
    last_name: str
    father_name: str
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    mobile: str
    personal_email: EmailStr
    pan: str
    aadhar: str
    uan: str
    pf_account_no: str
    bank_name: str
    account_no: str
    ifsc: str
    branch: str
    address: str
    pincode: str
    city: str
    monthly_gross_inr: float
    department: Literal['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']
    projects: List[str] = Field(default_factory=list)
    approver_user_id: str

class Approval(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"appr_{uuid.uuid4().hex[:8]}")
    item_type: Literal['client', 'contractor', 'employee']
    item_id: str
    requested_by: str
    status: Literal['Requested', 'Approved', 'Rejected', 'Hold'] = 'Requested'
    approved_by: Optional[str] = None
    approved_at: Optional[str] = None
    notes: Optional[str] = None
    staff_remarks: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ApprovalAction(BaseModel):
    action: Literal['approve', 'reject', 'hold']
    notes: Optional[str] = None

class ApprovalRequest(BaseModel):
    staff_remarks: Optional[str] = None

class SLAGenerateRequest(BaseModel):
    client_name: str
    address: str
    start_date: str
    tenure_months: int
    currency_preference: Literal['USD', 'INR']
    service: Literal['PPC', 'SEO', 'Content', 'Backlink']
    amount_ppc: Optional[float] = None
    amount_seo: Optional[float] = None
    amount: Optional[float] = None
    authorised_signatory: str
    designation: str

class NDAGenerateRequest(BaseModel):
    client_name: str
    address: str
    start_date: str
    authorised_signatory: str
    designation: str

class ICAGenerateRequest(BaseModel):
    contractor_name: str
    address: str
    start_date: str
    tenure_months: int
    amount_inr: float
    designation: str

class OfferLetterGenerateRequest(BaseModel):
    employee_name: str
    date: str
    gross_salary_lpa: float
    sign_before_date: str
    position: str
    department: str

class Asset(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"asset_{uuid.uuid4().hex[:8]}")
    asset_type: str
    model: str
    serial_number: str
    purchase_date: str
    vendor: str
    value_ex_gst: float
    warranty_period_months: int
    alloted_to: str
    email: EmailStr
    department: Literal['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']
    warranty_status: str = "Active"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class AssetCreate(BaseModel):
    asset_type: str
    model: str
    serial_number: str
    purchase_date: str
    vendor: str
    value_ex_gst: float
    warranty_period_months: int
    alloted_to: str
    email: EmailStr
    department: Literal['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']

# ============= HELPER FUNCTIONS =============

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        'user_id': user_id,
        'email': email,
        'role': role,
        'exp': datetime.now(timezone.utc) + timedelta(hours=24)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

def calculate_end_date(start_date: str, tenure_months: int) -> str:
    from dateutil.relativedelta import relativedelta
    start = datetime.fromisoformat(start_date)
    end = start + relativedelta(months=tenure_months)
    return end.isoformat()[:10]

def check_agreement_status(end_date: str) -> str:
    end = datetime.fromisoformat(end_date)
    # Remove timezone info for comparison
    today = datetime.now().date()
    end_date_only = end.date() if hasattr(end, 'date') else end
    return 'Live' if today <= end_date_only else 'Expired'

# ============= AUTH ROUTES =============

@api_router.post("/auth/login")
async def login(request: LoginRequest):
    user = await db.users.find_one({"email": request.email})
    if not user or not verify_password(request.password, user.get('password_hash', '')):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Generate OTP (for MVP, we'll use a simple 6-digit code)
    otp = str(random.randint(100000, 999999))
    await db.otps.update_one(
        {"email": request.email},
        {"$set": {"otp": otp, "created_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    
    return {
        "message": "OTP sent to email",
        "email": request.email,
        "otp": otp,  # For MVP, returning OTP (in production, send via email)
        "requires_verification": not user.get('otp_verified', False)
    }

@api_router.post("/auth/verify-otp")
async def verify_otp(request: OTPVerifyRequest):
    otp_record = await db.otps.find_one({"email": request.email})
    if not otp_record or otp_record['otp'] != request.otp:
        raise HTTPException(status_code=400, detail="Invalid OTP")
    
    user = await db.users.find_one({"email": request.email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Mark user as verified
    await db.users.update_one(
        {"email": request.email},
        {"$set": {"otp_verified": True}}
    )
    
    # Generate JWT token
    token = create_token(user['id'], user['email'], user['role'])
    
    return {
        "token": token,
        "user": {
            "id": user['id'],
            "name": user['name'],
            "email": user['email'],
            "role": user['role']
        }
    }

@api_router.get("/auth/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    user = await db.users.find_one({"id": current_user['user_id']}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

# ============= USER ROUTES =============

@api_router.get("/users", response_model=List[User])
async def get_users(current_user: dict = Depends(get_current_user)):
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).to_list(1000)
    return users

@api_router.post("/users", response_model=User)
async def create_user(user_data: UserCreate, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can create users")
    
    # Check if user exists
    existing = await db.users.find_one({"email": user_data.email})
    if existing:
        raise HTTPException(status_code=400, detail="User already exists")
    
    user = User(
        name=user_data.name,
        email=user_data.email,
        mobile=user_data.mobile,
        role=user_data.role,
        password_hash=hash_password(user_data.password),
        status='Active'
    )
    
    doc = user.model_dump()
    await db.users.insert_one(doc)
    return user

@api_router.patch("/users/{user_id}")
async def update_user(user_id: str, role: str = None, status: str = None, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can update users")
    
    # Check if user being updated is admin
    user_to_update = await db.users.find_one({"id": user_id})
    if user_to_update and user_to_update.get('role') == 'Admin' and role and role != 'Admin':
        raise HTTPException(status_code=403, detail="Cannot change Admin role")
    
    update_data = {}
    if role:
        update_data['role'] = role
    if status:
        update_data['status'] = status
    
    await db.users.update_one({"id": user_id}, {"$set": update_data})
    return {"message": "User updated successfully"}

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can delete users")
    
    # Check if user being deleted is admin
    user_to_delete = await db.users.find_one({"id": user_id})
    if not user_to_delete:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user_to_delete.get('role') == 'Admin':
        raise HTTPException(status_code=403, detail="Cannot delete Admin user")
    
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "User deleted successfully"}


# ============= CLIENT ROUTES =============

@api_router.get("/clients", response_model=List[Client])
async def get_clients(
    current_user: dict = Depends(get_current_user),
    sort_by: str = None,
    sort_order: str = 'asc',
    filter_status: str = None,
    filter_department: str = None
):
    query = {}
    if filter_status:
        query['client_status'] = filter_status
    if filter_department:
        query['service'] = filter_department
    
    clients = await db.clients.find(query, {"_id": 0}).to_list(1000)
    
    # Sorting
    if sort_by:
        reverse = sort_order == 'desc'
        clients.sort(key=lambda x: x.get(sort_by, ''), reverse=reverse)
    
    return clients

@api_router.get("/clients/active-by-department")
async def get_active_clients_by_department(
    department: str = None,
    current_user: dict = Depends(get_current_user)
):
    """Get active clients filtered by service/department for project assignment"""
    query = {"client_status": "Active"}
    if department:
        query['service'] = department
    
    clients = await db.clients.find(query, {"_id": 0, "id": 1, "client_name": 1, "service": 1}).to_list(1000)
    return clients


@api_router.post("/clients", response_model=Client)
async def create_client(client_data: ClientCreate, current_user: dict = Depends(get_current_user)):
    client = Client(**client_data.model_dump())
    client.end_date = calculate_end_date(client.start_date, client.tenure_months)
    client.agreement_status = check_agreement_status(client.end_date)
    
    doc = client.model_dump()
    await db.clients.insert_one(doc)
    return client

@api_router.patch("/clients/{client_id}")
async def update_client(client_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    # Recalculate end_date and agreement_status if start_date or tenure_months changed
    if 'start_date' in update_data or 'tenure_months' in update_data:
        client = await db.clients.find_one({"id": client_id})
        if client:
            start_date = update_data.get('start_date', client.get('start_date'))
            tenure_months = update_data.get('tenure_months', client.get('tenure_months'))
            
            if start_date and tenure_months:
                end_date = calculate_end_date(start_date, tenure_months)
                agreement_status = check_agreement_status(end_date)
                update_data['end_date'] = end_date
                update_data['agreement_status'] = agreement_status
    
    await db.clients.update_one({"id": client_id}, {"$set": update_data})
    return {"message": "Client updated successfully"}

@api_router.delete("/clients/{client_id}")
async def delete_client(client_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a client - Admin and Director only"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete clients")
    
    result = await db.clients.delete_one({"id": client_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")
    
    return {"message": "Client deleted successfully"}

@api_router.post("/clients/generate-sla")
async def generate_sla(request: SLAGenerateRequest):
    try:
        # Try template-based generation first
        if request.service == 'PPC':
            template_path = ROOT_DIR / 'templates' / 'SLA_PPC.docx'
        elif request.service == 'SEO':
            template_path = ROOT_DIR / 'templates' / 'SLA_SEO.docx'
        else:
            template_path = ROOT_DIR / 'templates' / 'SLA_PPC.docx'
        
        if template_path.exists():
            try:
                output_path = f"/tmp/SLA_{request.client_name.replace(' ', '_')}_{uuid.uuid4().hex[:6]}.docx"
                shutil.copy(template_path, output_path)
                
                merge_data = {
                    'client_name': request.client_name,
                    'address': request.address,
                    'start_date': request.start_date,
                    'tenure_months': str(request.tenure_months),
                    'service': request.service,
                    'currency': request.currency_preference,
                    'authorised_signatory': request.authorised_signatory,
                    'designation': request.designation,
                }
                
                if request.service == 'Both':
                    merge_data['amount_ppc'] = str(request.amount_ppc) if request.amount_ppc else '0'
                    merge_data['amount_seo'] = str(request.amount_seo) if request.amount_seo else '0'
                    merge_data['amount'] = str((request.amount_ppc or 0) + (request.amount_seo or 0))
                else:
                    merge_data['amount'] = str(request.amount) if request.amount else '0'
                
                document = MailMerge(output_path)
                document.merge(**merge_data)
                document.write(output_path)
                
                with open(output_path, 'rb') as f:
                    content = f.read()
                
                return Response(
                    content=content,
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': f'attachment; filename="SLA_{request.client_name.replace(" ", "_")}.docx"'}
                )
            except Exception as template_error:
                logger.error(f"Template merge error: {str(template_error)}")
                # Fall through to simple generation
        
    except Exception as e:
        logger.error(f"SLA generation error: {str(e)}")
    
    # Fallback: Generate simple document
    doc = Document()
    
    # Add header with logo placeholder
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('SERVICE LEVEL AGREEMENT')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1  # Center
    
    doc.add_paragraph()  # Spacing
    
    # Agreement details
    doc.add_paragraph(f"This Service Level Agreement (\"SLA\") is entered into on {request.start_date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Client Name: {request.client_name}")
    doc.add_paragraph(f"Address: {request.address}")
    doc.add_paragraph(f"Service Type: {request.service}")
    
    if request.service == 'Both':
        doc.add_paragraph(f"PPC Service Fee: {request.currency_preference} {request.amount_ppc:,.2f}")
        doc.add_paragraph(f"SEO Service Fee: {request.currency_preference} {request.amount_seo:,.2f}")
        doc.add_paragraph(f"Total Monthly Fee: {request.currency_preference} {(request.amount_ppc + request.amount_seo):,.2f}")
    else:
        doc.add_paragraph(f"Monthly Service Fee: {request.currency_preference} {request.amount:,.2f}")
    
    doc.add_paragraph(f"Contract Period: {request.tenure_months} months")
    doc.add_paragraph()
    
    # Signature section
    doc.add_paragraph("For and on behalf of the Client:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.authorised_signatory}")
    doc.add_paragraph(f"Designation: {request.designation}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="SLA_{request.client_name.replace(" ", "_")}.docx"'}
    )

@api_router.post("/clients/generate-nda")
async def generate_nda(request: NDAGenerateRequest):
    # Generate simple NDA document
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('NON-DISCLOSURE AGREEMENT')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1
    
    doc.add_paragraph()
    
    doc.add_paragraph(f"This Non-Disclosure Agreement (\"NDA\") is entered into on {request.start_date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Client Name: {request.client_name}")
    doc.add_paragraph(f"Address: {request.address}")
    doc.add_paragraph()
    
    doc.add_paragraph("This agreement governs the disclosure of confidential information between the parties.")
    doc.add_paragraph()
    
    # Signature section
    doc.add_paragraph("For and on behalf of the Client:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.authorised_signatory}")
    doc.add_paragraph(f"Designation: {request.designation}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="NDA_{request.client_name.replace(" ", "_")}.docx"'}
    )

# ============= CONTRACTOR ROUTES =============

@api_router.get("/contractors", response_model=List[Contractor])
async def get_contractors(
    current_user: dict = Depends(get_current_user),
    sort_by: str = None,
    sort_order: str = 'asc',
    filter_status: str = None,
    filter_department: str = None
):
    query = {}
    if filter_status:
        query['status'] = filter_status
    if filter_department:
        query['department'] = filter_department
    
    contractors = await db.contractors.find(query, {"_id": 0}).to_list(1000)
    
    # Sorting
    if sort_by:
        reverse = sort_order == 'desc'
        contractors.sort(key=lambda x: x.get(sort_by, ''), reverse=reverse)
    
    return contractors

@api_router.post("/contractors", response_model=Contractor)
async def create_contractor(contractor_data: ContractorCreate, current_user: dict = Depends(get_current_user)):
    contractor = Contractor(**contractor_data.model_dump())
    contractor.end_date = calculate_end_date(contractor.start_date, contractor.tenure_months)
    contractor.agreement_status = check_agreement_status(contractor.end_date)
    
    doc = contractor.model_dump()
    await db.contractors.insert_one(doc)
    return contractor

@api_router.patch("/contractors/{contractor_id}")
async def update_contractor(contractor_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    # Recalculate end_date and agreement_status if start_date or tenure_months changed
    if 'start_date' in update_data or 'tenure_months' in update_data:
        contractor = await db.contractors.find_one({"id": contractor_id})
        if contractor:
            start_date = update_data.get('start_date', contractor.get('start_date'))
            tenure_months = update_data.get('tenure_months', contractor.get('tenure_months'))
            
            if start_date and tenure_months:
                end_date = calculate_end_date(start_date, tenure_months)
                agreement_status = check_agreement_status(end_date)
                update_data['end_date'] = end_date
                update_data['agreement_status'] = agreement_status
    
    await db.contractors.update_one({"id": contractor_id}, {"$set": update_data})
    return {"message": "Contractor updated successfully"}

@api_router.delete("/contractors/{contractor_id}")
async def delete_contractor(contractor_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a contractor - Admin and Director only"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete contractors")
    
    result = await db.contractors.delete_one({"id": contractor_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Contractor not found")
    
    return {"message": "Contractor deleted successfully"}

@api_router.post("/contractors/generate-ica")
async def generate_ica(request: ICAGenerateRequest):
    # Generate simple ICA document
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('INDEPENDENT CONTRACTOR AGREEMENT')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1
    
    doc.add_paragraph()
    
    doc.add_paragraph(f"This Independent Contractor Agreement is entered into on {request.start_date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Contractor Name: {request.contractor_name}")
    doc.add_paragraph(f"Address: {request.address}")
    doc.add_paragraph(f"Designation: {request.designation}")
    doc.add_paragraph(f"Monthly Retainer: INR {request.amount_inr:,.2f}")
    doc.add_paragraph(f"Contract Period: {request.tenure_months} months")
    doc.add_paragraph()
    
    doc.add_paragraph("Terms and Conditions:")
    doc.add_paragraph("1. The Contractor agrees to provide services as per the scope of work.")
    doc.add_paragraph("2. Payment shall be made on a monthly basis.")
    doc.add_paragraph("3. This agreement can be terminated by either party with 30 days notice.")
    doc.add_paragraph()
    
    # Signature section
    doc.add_paragraph("Contractor Signature:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.contractor_name}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="ICA_{request.contractor_name.replace(" ", "_")}.docx"'}
    )

# ============= EMPLOYEE ROUTES =============

@api_router.get("/employees", response_model=List[Employee])
async def get_employees(
    current_user: dict = Depends(get_current_user),
    sort_by: str = None,
    sort_order: str = 'asc',
    filter_status: str = None,
    filter_department: str = None
):
    query = {}
    if filter_status:
        query['status'] = filter_status
    if filter_department:
        query['department'] = filter_department
    
    employees = await db.employees.find(query, {"_id": 0}).to_list(1000)
    
    # Sorting
    if sort_by:
        reverse = sort_order == 'desc'
        employees.sort(key=lambda x: x.get(sort_by, ''), reverse=reverse)
    
    return employees

@api_router.post("/employees", response_model=Employee)
async def create_employee(employee_data: EmployeeCreate, current_user: dict = Depends(get_current_user)):
    employee = Employee(**employee_data.model_dump())
    
    doc = employee.model_dump()
    await db.employees.insert_one(doc)
    return employee

@api_router.patch("/employees/{employee_id}")
async def update_employee(employee_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    await db.employees.update_one({"id": employee_id}, {"$set": update_data})
    return {"message": "Employee updated successfully"}

@api_router.delete("/employees/{employee_id}")
async def delete_employee(employee_id: str, current_user: dict = Depends(get_current_user)):
    """Delete an employee - Admin and Director only"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete employees")
    
    result = await db.employees.delete_one({"id": employee_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Employee not found")
    
    return {"message": "Employee deleted successfully"}

@api_router.post("/employees/generate-offer")
async def generate_offer_letter(request: OfferLetterGenerateRequest):
    # Calculate CTC
    gross_annual = request.gross_salary_lpa * 100000
    ctc_annual = gross_annual + 21600
    monthly_ctc = ctc_annual / 12
    monthly_gross = gross_annual / 12
    
    # Generate offer letter document
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('OFFER LETTER')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1
    
    doc.add_paragraph()
    doc.add_paragraph(f"Date: {request.date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Dear {request.employee_name},")
    doc.add_paragraph()
    
    doc.add_paragraph(f"We are pleased to offer you the position of {request.position} in the {request.department} department.")
    doc.add_paragraph()
    
    doc.add_paragraph("Compensation Details:")
    doc.add_paragraph(f"• Gross Annual Salary: INR {gross_annual:,.2f}")
    doc.add_paragraph(f"• Cost to Company (Annual): INR {ctc_annual:,.2f}")
    doc.add_paragraph(f"• Monthly CTC: INR {monthly_ctc:,.2f}")
    doc.add_paragraph(f"• Monthly Gross: INR {monthly_gross:,.2f}")
    doc.add_paragraph()
    
    # Salary breakdown table
    doc.add_paragraph("Monthly Salary Breakdown:")
    
    # Calculate components
    basic = monthly_gross * 0.50
    hra = monthly_gross * 0.30
    special = monthly_gross * 0.20
    
    doc.add_paragraph(f"• Basic Salary: INR {basic:,.2f}")
    doc.add_paragraph(f"• HRA: INR {hra:,.2f}")
    doc.add_paragraph(f"• Special Allowance: INR {special:,.2f}")
    doc.add_paragraph(f"• Employer PF Contribution: INR 1,800.00")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Please sign and return this offer letter before {request.sign_before_date}.")
    doc.add_paragraph()
    
    doc.add_paragraph("We look forward to welcoming you to our team!")
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Piperocket HR Team")
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph("Employee Acceptance:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.employee_name}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="Offer_{request.employee_name.replace(" ", "_")}.docx"'}
    )

# ============= APPROVAL ROUTES =============

@api_router.get("/approvals", response_model=List[Approval])
async def get_approvals(current_user: dict = Depends(get_current_user)):
    approvals = await db.approvals.find({}, {"_id": 0}).to_list(1000)
    return approvals

@api_router.post("/approvals/{item_type}/{item_id}/request")
async def request_approval(item_type: str, item_id: str, request: ApprovalRequest, current_user: dict = Depends(get_current_user)):
    if current_user['role'] == 'Director':
        raise HTTPException(status_code=403, detail="Directors cannot request approval")
    
    approval = Approval(
        item_type=item_type,
        item_id=item_id,
        requested_by=current_user['user_id'],
        status='Requested',
        staff_remarks=request.staff_remarks
    )
    
    doc = approval.model_dump()
    await db.approvals.insert_one(doc)
    return approval

@api_router.post("/approvals/{approval_id}/action")
async def approval_action(approval_id: str, action: ApprovalAction, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Director':
        raise HTTPException(status_code=403, detail="Only Directors can approve/reject/hold")
    
    status_map = {'approve': 'Approved', 'reject': 'Rejected', 'hold': 'Hold'}
    status = status_map.get(action.action, 'Requested')
    
    await db.approvals.update_one(
        {"id": approval_id},
        {"$set": {
            "status": status,
            "approved_by": current_user['user_id'],
            "approved_at": datetime.now(timezone.utc).isoformat(),
            "notes": action.notes
        }}
    )
    
    return {"message": f"Approval {status.lower()} successfully"}

@api_router.delete("/approvals/reset")
async def reset_approvals(current_user: dict = Depends(get_current_user)):
    """Reset all approval records - accessible by Staff and Admin"""
    if current_user['role'] == 'Director':
        raise HTTPException(status_code=403, detail="Directors cannot reset approvals")
    
    result = await db.approvals.delete_many({})
    return {"message": f"Reset complete. Deleted {result.deleted_count} approval records"}

# ============= DASHBOARD ROUTES =============

@api_router.get("/dashboard/summary")
async def get_dashboard_summary(current_user: dict = Depends(get_current_user)):
    # Get alerts
    today = datetime.now(timezone.utc)
    thirty_days_later = today + timedelta(days=30)
    
    # Expiring agreements - get actual client names
    clients_all = await db.clients.find({"client_status": "Active"}).to_list(1000)
    expiring_clients = []
    for client in clients_all:
        if client.get('end_date'):
            try:
                end_date = datetime.fromisoformat(client['end_date'])
                # Remove timezone for comparison
                if hasattr(end_date, 'date'):
                    end_date_only = end_date.date()
                else:
                    end_date_only = end_date
                
                if today.date() <= end_date_only <= thirty_days_later.date():
                    expiring_clients.append({
                        "name": client['client_name'],
                        "end_date": client['end_date'],
                        "service": client['service']
                    })
            except:
                pass
    
    # Upcoming birthdays - get employee and contractor names
    employees = await db.employees.find({"status": "Active"}).to_list(1000)
    contractors = await db.contractors.find({"status": "Active"}).to_list(1000)
    
    upcoming_birthdays = []
    
    for emp in employees:
        if emp.get('dob'):
            try:
                dob = datetime.fromisoformat(emp['dob'])
                # Calculate days until birthday this year
                this_year_bday = datetime(today.year, dob.month, dob.day)
                if this_year_bday < datetime.now():
                    # Birthday already passed this year, check next year
                    this_year_bday = datetime(today.year + 1, dob.month, dob.day)
                
                days_until = (this_year_bday - datetime.now()).days
                if 0 <= days_until <= 15:
                    upcoming_birthdays.append({
                        "name": f"{emp['first_name']} {emp['last_name']}",
                        "date": emp['dob'],
                        "type": "Employee",
                        "department": emp.get('department', '')
                    })
            except:
                pass
    
    for con in contractors:
        if con.get('dob'):
            try:
                dob = datetime.fromisoformat(con['dob'])
                this_year_bday = datetime(today.year, dob.month, dob.day)
                if this_year_bday < datetime.now():
                    this_year_bday = datetime(today.year + 1, dob.month, dob.day)
                
                days_until = (this_year_bday - datetime.now()).days
                if 0 <= days_until <= 15:
                    upcoming_birthdays.append({
                        "name": con['name'],
                        "date": con['dob'],
                        "type": "Contractor",
                        "department": con.get('department', '')
                    })
            except:
                pass
    
    # Sort birthdays by date
    upcoming_birthdays.sort(key=lambda x: datetime.fromisoformat(x['date']))
    
    # Expired agreements
    expired_clients = []
    for client in clients_all:
        if client.get('end_date'):
            try:
                end_date = datetime.fromisoformat(client['end_date'])
                if hasattr(end_date, 'date'):
                    end_date_only = end_date.date()
                else:
                    end_date_only = end_date
                
                if end_date_only < today.date():
                    expired_clients.append({
                        "name": client['client_name'],
                        "end_date": client['end_date'],
                        "service": client['service']
                    })
            except:
                pass
    
    # Revenue metrics
    clients = await db.clients.find({"client_status": "Active"}).to_list(1000)
    revenue_by_dept = {}
    for dept in ['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']:
        dept_clients = [c for c in clients if c.get('service') == dept]
        count = len(dept_clients)
        amount = sum(c.get('amount_inr', 0) for c in dept_clients)
        revenue_by_dept[dept] = {"count": count, "amount": amount}
    
    # Employee metrics
    employee_by_dept = {}
    for dept in ['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']:
        dept_employees = [e for e in employees if e.get('department') == dept]
        count = len(dept_employees)
        cost = sum(e.get('monthly_gross_inr', 0) for e in dept_employees)
        employee_by_dept[dept] = {"count": count, "cost": cost}
    
    # Contractor metrics
    contractor_by_dept = {}
    for dept in ['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']:
        dept_contractors = [c for c in contractors if c.get('department') == dept]
        count = len(dept_contractors)
        cost = sum(c.get('monthly_retainer_inr', 0) for c in dept_contractors)
        contractor_by_dept[dept] = {"count": count, "cost": cost}
    
    return {
        "alerts": {
            "expiring_agreements": expiring_clients,
            "expired_agreements": expired_clients,
            "upcoming_birthdays": upcoming_birthdays
        },
        "revenue": revenue_by_dept,
        "employees": employee_by_dept,
        "contractors": contractor_by_dept
    }

# ============= BULK EXPORT/IMPORT ROUTES =============

@api_router.get("/clients/export")
async def export_clients(current_user: dict = Depends(get_current_user)):
    """Export all clients to Excel"""
    clients = await db.clients.find({}, {"_id": 0}).to_list(1000)
    
    if not clients:
        raise HTTPException(status_code=404, detail="No clients to export")
    
    df = pd.DataFrame(clients)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Clients')
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="clients_export.xlsx"'}
    )

@api_router.get("/contractors/export")
async def export_contractors(current_user: dict = Depends(get_current_user)):
    """Export all contractors to Excel"""
    contractors = await db.contractors.find({}, {"_id": 0}).to_list(1000)
    
    if not contractors:
        raise HTTPException(status_code=404, detail="No contractors to export")
    
    df = pd.DataFrame(contractors)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Contractors')
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="contractors_export.xlsx"'}
    )

@api_router.get("/employees/export")
async def export_employees(current_user: dict = Depends(get_current_user)):
    """Export all employees to Excel"""
    employees = await db.employees.find({}, {"_id": 0}).to_list(1000)
    
    if not employees:
        raise HTTPException(status_code=404, detail="No employees to export")
    
    df = pd.DataFrame(employees)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Employees')
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="employees_export.xlsx"'}
    )

@api_router.get("/clients/sample")
async def get_client_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Clients"
    
    # Headers with bold formatting
    headers = ['client_name', 'address', 'start_date', 'tenure_months', 'currency_preference', 
               'service', 'amount_inr', 'authorised_signatory', 'signatory_designation', 
               'gst', 'poc_name', 'poc_email', 'poc_designation', 'poc_mobile', 'approver_user_id']
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    # Sample rows
    ws.append(['ABC Corp', '123 Main St', '2025-01-01', 12, 'INR', 'PPC', 50000, 
               'John Doe', 'CEO', 'GST123', 'Jane Smith', 'jane@abc.com', 'Manager', '9876543210', 'user_id'])
    ws.append(['XYZ Ltd', '456 Park Ave', '2025-02-01', 6, 'INR', 'SEO', 75000, 
               'Mike Johnson', 'Director', 'GST456', 'Sarah Lee', 'sarah@xyz.com', 'Lead', '9876543211', 'user_id'])
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="client_sample.xlsx"'}
    )

@api_router.get("/contractors/sample")
async def get_contractor_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Contractors"
    
    headers = ['name', 'doj', 'start_date', 'tenure_months', 'dob', 'pan', 'aadhar', 
               'mobile', 'personal_email', 'bank_name', 'account_holder', 'account_no', 
               'ifsc', 'address_1', 'pincode', 'city', 'address_2', 'department', 
               'monthly_retainer_inr', 'designation', 'approver_user_id']
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    ws.append(['John Contractor', '2025-01-01', '2025-01-01', 6, '1990-05-15', 'ABCDE1234F', 
               '123456789012', '9876543210', 'john@email.com', 'Bank Name', 'John Contractor', 
               '1234567890', 'BANK0001234', '123 Street', '110001', 'Delhi', 'Near Market', 
               'PPC', 35000, 'Consultant', 'user_id'])
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="contractor_sample.xlsx"'}
    )

@api_router.get("/employees/sample")
async def get_employee_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Employees"
    
    headers = ['doj', 'work_email', 'emp_id', 'first_name', 'last_name', 'father_name', 
               'dob', 'mobile', 'personal_email', 'pan', 'aadhar', 'uan', 'pf_account_no', 
               'bank_name', 'account_no', 'ifsc', 'branch', 'address', 'pincode', 'city', 
               'monthly_gross_inr', 'department', 'approver_user_id']
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    ws.append(['2025-01-15', 'john@company.com', 'EMP001', 'John', 'Doe', 'James Doe', 
               '1995-03-20', '9876543210', 'john.personal@email.com', 'ABCDE1234F', 
               '123456789012', 'UAN123456', 'PF123456', 'Bank Name', '1234567890', 
               'BANK0001234', 'Main Branch', '123 Street', '110001', 'Delhi', 
               60000, 'PPC', 'user_id'])
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="employee_sample.xlsx"'}
    )

@api_router.get("/assets/sample")
async def get_asset_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Assets"
    
    headers = ['asset_type', 'model', 'serial_number', 'purchase_date', 'vendor', 
               'value_ex_gst', 'warranty_period_months', 'alloted_to', 'email', 'department']
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    ws.append(['Laptop', 'Dell XPS 15', 'SN123456', '2024-01-15', 'Dell India', 
               75000, 12, 'John Doe', 'john@company.com', 'PPC'])
    ws.append(['Monitor', 'LG 27inch', 'SN789012', '2024-02-01', 'LG Store', 
               15000, 24, 'Jane Smith', 'jane@company.com', 'SEO'])
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="asset_sample.xlsx"'}
    )

@api_router.post("/clients/import")
async def import_clients(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import clients from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                # Create client object
                client_data = ClientCreate(
                    client_name=str(row['client_name']),
                    address=str(row['address']),
                    start_date=str(row['start_date'])[:10],
                    tenure_months=int(row['tenure_months']),
                    currency_preference=str(row.get('currency_preference', 'INR')),
                    service=str(row['service']),
                    amount_inr=float(row['amount_inr']),
                    authorised_signatory=str(row['authorised_signatory']),
                    signatory_designation=str(row['signatory_designation']),
                    gst=str(row['gst']),
                    poc_name=str(row['poc_name']),
                    poc_email=str(row['poc_email']),
                    poc_designation=str(row['poc_designation']),
                    poc_mobile=str(row['poc_mobile']),
                    approver_user_id=str(row['approver_user_id'])
                )
                
                client = Client(**client_data.model_dump())
                client.end_date = calculate_end_date(client.start_date, client.tenure_months)
                client.agreement_status = check_agreement_status(client.end_date)
                
                await db.clients.insert_one(client.model_dump())
                imported_count += 1
                
            except Exception as e:
                errors.append(f"Row {index + 2}: {str(e)}")
        
        return {
            "message": f"Import completed. {imported_count} clients imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

@api_router.post("/contractors/import")
async def import_contractors(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import contractors from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                contractor_data = ContractorCreate(
                    name=str(row['name']),
                    doj=str(row['doj'])[:10],
                    start_date=str(row['start_date'])[:10],
                    tenure_months=int(row['tenure_months']),
                    dob=str(row['dob'])[:10],
                    pan=str(row['pan']),
                    aadhar=str(row['aadhar']),
                    mobile=str(row['mobile']),
                    personal_email=str(row['personal_email']),
                    bank_name=str(row['bank_name']),
                    account_holder=str(row['account_holder']),
                    account_no=str(row['account_no']),
                    ifsc=str(row['ifsc']),
                    address_1=str(row['address_1']),
                    pincode=str(row['pincode']),
                    city=str(row['city']),
                    address_2=str(row.get('address_2', '')),
                    department=str(row['department']),
                    monthly_retainer_inr=float(row['monthly_retainer_inr']),
                    designation=str(row['designation']),
                    approver_user_id=str(row['approver_user_id'])
                )
                
                contractor = Contractor(**contractor_data.model_dump())
                contractor.end_date = calculate_end_date(contractor.start_date, contractor.tenure_months)
                contractor.agreement_status = check_agreement_status(contractor.end_date)
                
                await db.contractors.insert_one(contractor.model_dump())
                imported_count += 1
                
            except Exception as e:
                errors.append(f"Row {index + 2}: {str(e)}")
        
        return {
            "message": f"Import completed. {imported_count} contractors imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

@api_router.post("/employees/import")
async def import_employees(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import employees from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                employee_data = EmployeeCreate(
                    doj=str(row['doj'])[:10],
                    work_email=str(row['work_email']),
                    emp_id=str(row['emp_id']),
                    first_name=str(row['first_name']),
                    last_name=str(row['last_name']),
                    father_name=str(row['father_name']),
                    dob=str(row['dob'])[:10],
                    mobile=str(row['mobile']),
                    personal_email=str(row['personal_email']),
                    pan=str(row['pan']),
                    aadhar=str(row['aadhar']),
                    uan=str(row['uan']),
                    pf_account_no=str(row['pf_account_no']),
                    bank_name=str(row['bank_name']),
                    account_no=str(row['account_no']),
                    ifsc=str(row['ifsc']),
                    branch=str(row['branch']),
                    address=str(row['address']),
                    pincode=str(row['pincode']),
                    city=str(row['city']),
                    monthly_gross_inr=float(row['monthly_gross_inr']),
                    department=str(row['department']),
                    approver_user_id=str(row['approver_user_id'])
                )
                
                employee = Employee(**employee_data.model_dump())
                await db.employees.insert_one(employee.model_dump())
                imported_count += 1
                
            except Exception as e:
                errors.append(f"Row {index + 2}: {str(e)}")
        
        return {
            "message": f"Import completed. {imported_count} employees imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

@api_router.post("/assets/import")
async def import_assets(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import assets from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                asset_data = AssetCreate(
                    asset_type=str(row['asset_type']),
                    model=str(row['model']),
                    serial_number=str(row['serial_number']),
                    purchase_date=str(row['purchase_date'])[:10],
                    vendor=str(row['vendor']),
                    value_ex_gst=float(row['value_ex_gst']),
                    warranty_period_months=int(row['warranty_period_months']),
                    alloted_to=str(row['alloted_to']),
                    email=str(row['email']),
                    department=str(row['department'])
                )
                
                asset = Asset(**asset_data.model_dump())
                
                # Calculate warranty status
                purchase_date = datetime.fromisoformat(asset.purchase_date)
                warranty_end = purchase_date + timedelta(days=asset.warranty_period_months * 30)
                asset.warranty_status = 'Active' if datetime.now() <= warranty_end else 'Expired'
                
                await db.assets.insert_one(asset.model_dump())
                imported_count += 1
                
            except Exception as e:
                errors.append(f"Row {index + 2}: {str(e)}")
        
        return {
            "message": f"Import completed. {imported_count} assets imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

# ============= ASSET TRACKER ROUTES =============

@api_router.get("/assets", response_model=List[Asset])
async def get_assets(
    current_user: dict = Depends(get_current_user),
    department: str = None
):
    query = {}
    if department:
        query['department'] = department
    
    assets = await db.assets.find(query, {"_id": 0}).to_list(1000)
    
    # Update warranty status
    for asset in assets:
        purchase_date = datetime.fromisoformat(asset['purchase_date'])
        warranty_end = purchase_date + timedelta(days=asset['warranty_period_months'] * 30)
        if datetime.now() > warranty_end:
            asset['warranty_status'] = 'Expired'
        else:
            asset['warranty_status'] = 'Active'
    
    return assets

@api_router.post("/assets", response_model=Asset)
async def create_asset(asset_data: AssetCreate, current_user: dict = Depends(get_current_user)):
    asset = Asset(**asset_data.model_dump())
    
    # Calculate warranty status
    purchase_date = datetime.fromisoformat(asset.purchase_date)
    warranty_end = purchase_date + timedelta(days=asset.warranty_period_months * 30)
    asset.warranty_status = 'Active' if datetime.now() <= warranty_end else 'Expired'
    
    doc = asset.model_dump()
    await db.assets.insert_one(doc)
    return asset

@api_router.patch("/assets/{asset_id}")
async def update_asset(asset_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    await db.assets.update_one({"id": asset_id}, {"$set": update_data})
    return {"message": "Asset updated successfully"}

@api_router.delete("/assets/{asset_id}")
async def delete_asset(asset_id: str, current_user: dict = Depends(get_current_user)):
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete assets")
    
    result = await db.assets.delete_one({"id": asset_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Asset not found")
    
    return {"message": "Asset deleted successfully"}

@api_router.get("/assets/export")
async def export_assets(current_user: dict = Depends(get_current_user)):
    """Export all assets to Excel"""
    assets = await db.assets.find({}, {"_id": 0}).to_list(1000)
    
    if not assets:
        raise HTTPException(status_code=404, detail="No assets to export")
    
    df = pd.DataFrame(assets)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Assets')
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="assets_export.xlsx"'}
    )

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def startup():
    # Create seed admin user
    existing = await db.users.find_one({"email": "Vishnu@onedotfinance.com"})
    if not existing:
        admin = User(
            name="Vishnu Admin",
            email="Vishnu@onedotfinance.com",
            role="Admin",
            password_hash=hash_password("12345678"),
            status="Active",
            otp_verified=False
        )
        await db.users.insert_one(admin.model_dump())
        logger.info("Admin user created")

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
